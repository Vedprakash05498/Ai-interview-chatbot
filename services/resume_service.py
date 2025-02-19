# -*- coding: utf-8 -*-
import re
from typing import BinaryIO, Dict, List
import PyPDF2
from docx import Document
import magic
import io
import os
import fitz  # PyMuPDF
import docx
from uuid import uuid4
from datetime import datetime

class ResumeService:
    ALLOWED_MIME_TYPES = [
        'application/pdf',
        'application/msword',  # .doc
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # .docx
    ]
    
    def __init__(self):
        self.nlp = None
        self.web_dev_keywords = {
            'html', 'css', 'javascript', 'react', 'angular', 'vue', 'node', 'frontend', 'backend'
        }
        self.ai_ml_keywords = {
            'machine learning', 'ai', 'deep learning', 'neural networks', 'python', 'tensorflow', 'pytorch'
        }
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            print(f"Warning: spaCy not available, using basic parsing: {str(e)}")
        self.upload_dir = "uploads/resumes"
        os.makedirs(self.upload_dir, exist_ok=True)

    def validate_file(self, file_content: bytes) -> bool:
        mime = magic.Magic(mime=True)
        file_type = mime.from_buffer(file_content)
        return file_type in self.ALLOWED_MIME_TYPES

    def save_resume(self, file_content: bytes, original_filename: str) -> str:
        if not self.validate_file(file_content):
            raise ValueError("Invalid file format. Only PDF, DOC, and DOCX files are allowed.")
        
        # Generate unique filename
        ext = os.path.splitext(original_filename)[1]
        filename = f"{uuid4()}{ext}"
        filepath = os.path.join(self.upload_dir, filename)
        
        with open(filepath, "wb") as f:
            f.write(file_content)
        
        return filepath

    def parse_resume(self, filepath: str) -> Dict:
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(filepath)
        elif ext in ['.doc', '.docx']:
            return self._parse_docx(filepath)
        else:
            raise ValueError("Unsupported file format")

    def _parse_pdf(self, filepath: str) -> Dict:
        text = ""
        with fitz.open(filepath) as doc:
            for page in doc:
                text += page.get_text()
        
        return self._extract_information(text)

    def _parse_docx(self, filepath: str) -> Dict:
        doc = docx.Document(filepath)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return self._extract_information(text)

    def _extract_information(self, text: str) -> Dict:
        # This is a basic implementation - you might want to use more sophisticated
        # NLP techniques or regex patterns for better extraction
        return {
            "skills": self._extract_skills(text),
            "education": self._extract_education(text),
            "experience": self._extract_experience(text),
            "projects": self._extract_projects(text)
        }

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        # Define comprehensive skill sets
        skill_keywords = {
            'ai_ml': {
                'machine learning', 'deep learning', 'neural networks', 'ai', 'artificial intelligence',
                'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'opencv', 'computer vision',
                'nlp', 'natural language processing', 'data science', 'pandas', 'numpy',
                'data analysis', 'statistics', 'python', 'r', 'jupyter',
                'big data', 'hadoop', 'spark', 'data mining', 'data visualization',
                'reinforcement learning', 'cnn', 'rnn', 'lstm', 'transformers'
            },
            'web_dev': {
                'html', 'css', 'javascript', 'typescript', 'react', 'angular', 'vue',
                'node.js', 'express', 'django', 'flask', 'sql', 'mongodb', 'postgresql',
                'rest api', 'graphql', 'aws', 'docker', 'kubernetes', 'ci/cd',
                'git', 'webpack', 'npm', 'yarn', 'redux', 'bootstrap', 'sass',
                'frontend', 'backend', 'full stack', 'web development', 'responsive design'
            }
        }

        found_skills = set()
        text_lower = text.lower()
        
        # Use regex to find skills
        for domain, skills in skill_keywords.items():
            for skill in skills:
                if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                    found_skills.add(skill)
        
        # If spaCy is available, use it for better extraction
        if self.nlp:
            doc = self.nlp(text)
            # Look for technical terms and proper nouns
            for token in doc:
                if token.pos_ in ['PROPN', 'NOUN'] and len(token.text) > 2:
                    skill = token.text.lower()
                    if skill in skill_keywords['ai_ml'] or skill in skill_keywords['web_dev']:
                        found_skills.add(skill)
        
        return list(found_skills)

    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education information"""
        education_info = []
        education_text = self._get_section(text, "EDUCATION")
        
        if not education_text:
            return []
        
        # Common degree patterns
        degree_patterns = [
            r'(B\.?Tech|Bachelor of Technology)',
            r'(M\.?Tech|Master of Technology)',
            r'(B\.?E|Bachelor of Engineering)',
            r'(M\.?S|Master of Science)',
            r'(B\.?Sc|Bachelor of Science)',
            r'(Ph\.?D|Doctorate)',
            r'(MBA|Master of Business Administration)'
        ]
        
        # Split into different education entries
        entries = education_text.split('\n\n')
        
        for entry in entries:
            edu = {}
            
            # Try to extract degree
            for pattern in degree_patterns:
                match = re.search(pattern, entry, re.IGNORECASE)
                if match:
                    edu['degree'] = match.group(0)
                    break
            
            # Try to extract year
            year_match = re.search(r'20\d{2}', entry)
            if year_match:
                edu['year'] = year_match.group(0)
            
            # Try to extract institution
            if edu:  # Only add if we found a degree
                education_info.append(edu)
        
        return education_info

    def _extract_experience(self, text: str) -> List[Dict]:
        # Implement experience extraction logic
        return []

    def _extract_projects(self, text: str) -> List[Dict]:
        # Implement project extraction logic
        return []

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")

    def _split_into_sections(self, text: str) -> Dict[str, List[str]]:
        """Split text into sections"""
        sections = {}
        current_section = None
        current_content = []
        
        # Split text into lines and process each line
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Check if this is a section header
            upper_line = line.upper()
            if any(section in upper_line for section in ['SKILLS', 'EXPERIENCE', 'EDUCATION']):
                # Save previous section if exists
                if current_section and current_content:
                    sections[current_section] = current_content
                
                # Start new section
                current_section = next(s for s in ['SKILLS', 'EXPERIENCE', 'EDUCATION'] if s in upper_line)
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Add the last section
        if current_section and current_content:
            sections[current_section] = current_content
        
        # Process skills section (comma-separated)
        if 'SKILLS' in sections:
            skills_text = ' '.join(sections['SKILLS'])
            sections['SKILLS'] = [s.strip() for s in skills_text.split(',') if s.strip()]
        
        # Limit section sizes
        return {
            'SKILLS': sections.get('SKILLS', [])[:10],
            'EXPERIENCE': sections.get('EXPERIENCE', [])[:5],
            'EDUCATION': sections.get('EDUCATION', [])[:3]
        }

    def _get_section(self, text: str, section_name: str) -> str:
        """Extract section from text"""
        pattern = f"{section_name}.*?(?=\\n\\n|$)"
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            section = match.group(0)
            section = re.sub(f"{section_name}.*?\\n", '', section, flags=re.IGNORECASE)
            return section.strip()
        return ""

    def _basic_extract_skills(self, text):
        """Basic skill extraction without spaCy"""
        common_skills = ['python', 'java', 'javascript', 'html', 'css', 'sql', 'react', 'angular', 'node']
        found_skills = []
        for skill in common_skills:
            if re.search(r'\b' + skill + r'\b', text.lower()):
                found_skills.append(skill)
        return found_skills

    def _basic_extract_experience(self, text):
        """Basic experience extraction without spaCy"""
        experiences = []
        lines = text.split('\n')
        for line in lines:
            if any(word in line.lower() for word in ['work', 'job', 'position', 'experience']):
                experiences.append(line.strip())
        return experiences[:5]

    def _basic_extract_education(self, text):
        """Basic education extraction without spaCy"""
        education = []
        lines = text.split('\n')
        for line in lines:
            if any(word in line.lower() for word in ['degree', 'university', 'college', 'school']):
                education.append(line.strip())
        return education[:3]

    def determine_domain(self, skills: List[str]) -> str:
        """Determine candidate's domain based on skills"""
        ai_ml_score = 0
        web_dev_score = 0
        
        # Weight different skills
        skill_weights = {
            'ai_ml': {
                'machine learning': 3, 'deep learning': 3, 'neural networks': 3,
                'tensorflow': 2, 'pytorch': 2, 'data science': 2,
                'python': 1, 'statistics': 1, 'data analysis': 1
            },
            'web_dev': {
                'react': 3, 'angular': 3, 'node.js': 3,
                'javascript': 2, 'typescript': 2, 'html': 2,
                'css': 1, 'git': 1, 'rest api': 1
            }
        }
        
        for skill in skills:
            skill_lower = skill.lower()
            # Check AI/ML skills
            if skill_lower in skill_weights['ai_ml']:
                ai_ml_score += skill_weights['ai_ml'][skill_lower]
            elif skill_lower in self.ai_ml_keywords:
                ai_ml_score += 1
            
            # Check Web Dev skills
            if skill_lower in skill_weights['web_dev']:
                web_dev_score += skill_weights['web_dev'][skill_lower]
            elif skill_lower in self.web_dev_keywords:
                web_dev_score += 1
        
        return 'ai_ml' if ai_ml_score >= web_dev_score else 'web_dev'

    def determine_skill_level(self, experience: List[str], skills: List[str]) -> str:
        """Determine candidate's skill level"""
        # Calculate years of experience
        years = self._extract_years_experience(experience)
        
        # Calculate skill breadth
        skill_breadth = len(skills)
        
        # Calculate skill depth (looking for advanced terms)
        advanced_terms = {
            'ai_ml': {'deep learning', 'neural networks', 'machine learning', 'research'},
            'web_dev': {'architecture', 'scalability', 'microservices', 'optimization'}
        }
        
        skill_depth = sum(1 for skill in skills if any(
            term in skill.lower() for term in advanced_terms['ai_ml'] | advanced_terms['web_dev']
        ))
        
        # Determine level based on multiple factors
        if (years >= 5 or skill_breadth >= 10) and skill_depth >= 3:
            return 'expert'
        elif (years >= 2 or skill_breadth >= 5) and skill_depth >= 1:
            return 'intermediate'
        else:
            return 'beginner'

    def _extract_years_experience(self, experience: List[str]) -> int:
        """Extract total years of experience"""
        total_years = 0
        year_pattern = r'(\d+)[\s-]*(year|yr)'
        
        for exp in experience:
            matches = re.findall(year_pattern, exp.lower())
            for match in matches:
                try:
                    total_years += int(match[0])
                except ValueError:
                    continue
                
        return total_years 